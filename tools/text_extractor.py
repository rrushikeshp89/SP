"""
text_extractor.py — Resume Text Extraction (Layer 3 Tool)
Extracts clean plain text from PDF, DOCX, and TXT files.
Implements the fallback chain: pdfplumber → PyMuPDF → OCR placeholder.
"""

import re
import time
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ExtractionResult:
    success: bool = False
    raw_text: str = ""
    page_count: int | None = None
    method_used: str = "none"
    error_message: str | None = None
    processing_time_ms: float = 0.0


_MIN_TEXT_LEN = 50  # Minimum chars to consider extraction successful


def extract_text(file_path: str | Path, file_format: str) -> ExtractionResult:
    """
    Extract text from a resume file.

    Args:
        file_path: Absolute path to the file.
        file_format: One of "pdf", "docx", "txt".

    Returns:
        ExtractionResult with success flag and extracted text.
    """
    start = time.perf_counter()
    result = ExtractionResult()

    # Validate file exists
    path = Path(file_path)
    if not path.exists():
        result.error_message = f"File not found: {file_path}"
        result.processing_time_ms = _elapsed(start)
        return result

    fmt = file_format.lower().strip()

    try:
        if fmt == "txt":
            result = _extract_txt(path)
        elif fmt == "docx":
            result = _extract_docx(path)
        elif fmt == "pdf":
            result = _extract_pdf(path)
        else:
            result.error_message = f"Unsupported format: {fmt}"
    except Exception as exc:
        result.success = False
        result.error_message = f"Unexpected error: {exc}"

    result.processing_time_ms = _elapsed(start)
    return result


# ── TXT ──────────────────────────────────────────────

def _extract_txt(path: Path) -> ExtractionResult:
    """Read text file with encoding fallback."""
    for encoding in ("utf-8", "latin-1"):
        try:
            text = path.read_text(encoding=encoding)
            return ExtractionResult(
                success=True,
                raw_text=_clean(text),
                method_used="direct",
            )
        except UnicodeDecodeError:
            continue

    return ExtractionResult(
        error_message=f"Could not decode TXT file: {path.name}",
    )


# ── DOCX ─────────────────────────────────────────────

def _extract_docx(path: Path) -> ExtractionResult:
    """Extract text from DOCX via python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        return ExtractionResult(
            error_message="python-docx not installed",
        )

    try:
        doc = Document(str(path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        text = "\n".join(parts)
        return ExtractionResult(
            success=True,
            raw_text=_clean(text),
            method_used="python-docx",
        )
    except Exception as exc:
        return ExtractionResult(
            error_message=f"DOCX parsing error: {exc}",
        )


# ── PDF ──────────────────────────────────────────────

def _extract_pdf(path: Path) -> ExtractionResult:
    """PDF extraction with pdfplumber → PyMuPDF fallback chain."""

    # Attempt 1: pdfplumber
    result = _pdf_pdfplumber(path)
    if result.success and len(result.raw_text.strip()) >= _MIN_TEXT_LEN:
        return result

    # Attempt 2: PyMuPDF
    result = _pdf_pymupdf(path)
    if result.success and len(result.raw_text.strip()) >= _MIN_TEXT_LEN:
        return result

    # Attempt 3: OCR placeholder
    return ExtractionResult(
        error_message="All extraction methods failed (OCR not yet implemented)",
    )


def _pdf_pdfplumber(path: Path) -> ExtractionResult:
    """Extract with pdfplumber."""
    try:
        import pdfplumber  # type: ignore[import-untyped]
    except ImportError:
        return ExtractionResult(error_message="pdfplumber not installed")

    try:
        pages_text: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)

        full_text = "\n\n".join(pages_text)
        return ExtractionResult(
            success=True,
            raw_text=_clean(full_text),
            page_count=len(pages_text),
            method_used="pdfplumber",
        )
    except Exception as exc:
        return ExtractionResult(
            error_message=f"pdfplumber error: {exc}",
        )


def _pdf_pymupdf(path: Path) -> ExtractionResult:
    """Extract with PyMuPDF (fitz)."""
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError:
        return ExtractionResult(error_message="PyMuPDF not installed")

    try:
        doc = fitz.open(str(path))
        pages_text: list[str] = []
        for page in doc:
            pages_text.append(page.get_text("text"))
        doc.close()

        full_text = "\n\n".join(pages_text)
        return ExtractionResult(
            success=True,
            raw_text=_clean(full_text),
            page_count=len(pages_text),
            method_used="pymupdf",
        )
    except Exception as exc:
        return ExtractionResult(
            error_message=f"PyMuPDF error: {exc}",
        )


# ── Helpers ──────────────────────────────────────────

def _clean(text: str) -> str:
    """Collapse 3+ consecutive newlines into 2, strip edges."""
    text = text.strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _elapsed(start: float) -> float:
    """Return elapsed time in ms since start."""
    return round((time.perf_counter() - start) * 1000, 2)
